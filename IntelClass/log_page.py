import flet as ft
from datepicker.datepicker import DatePicker
from datepicker.selection_type import SelectionType
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL

class AttackReport:
    def __init__(self,attack_logs, logs_id_list):
        self.attack_logs = attack_logs
        self.logs_id_list = logs_id_list
        self.description = {
            "Worms":"Описание червей",
            "DDoS":"Произошла DDoS-атака, направленная на сетевые ресурсы, с использованием адресов 192.168.0.1 и 192.168.0.2.\nАтака имела характер ICMP флуда.",
            "Exploits":"Описание эксплойтов",
            "Analysis":"Описание анализа",
            "Backdoor":"Описание бэкдора",
            "Shellcode":"Описание шеллкода"
        }
        self.recommendations = {
            "Worms":["Рекомендация по червям 1","Рекомендация по червям 2","Рекомендация по червям 3"],
            "DDoS":["Добавить указанные источники в черный список сетевого экрана"],
            "Exploits":["Рекомендация по эксплойтам 1","Рекомендация по эксплойтам 2","Рекомендация по эксплойтам 3"],
            "Analysis":["Рекомендация по анализу 1","Рекомендация по анализу 2","Рекомендация по анализу 3"],
            "Backdoor":["Рекомендация по бэкдору 1","Рекомендация по бэкдору 2","Рекомендация по бэкдору 3"],
            "Shellcode":["Рекомендация по шеллкоду 1","Рекомендация по шеллкоду 2","Рекомендация по шеллкоду 3"]
        }
        self.document = Document()
        self.incidents = list(filter(self.id_filter, self.attack_logs))

    def id_filter(self, record):
        if record.id in self.logs_id_list:
            return True
        else:
            return False
            
    def add_attack_info(self):
        title_style = self.document.styles.add_style('Заголовок отчёта', WD_STYLE_TYPE.PARAGRAPH)
        title_paragraph_format = title_style.paragraph_format
        title_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(16)
        title_font.bold = True

        incident_title_style = self.document.styles.add_style('Заголовок инцидента', WD_STYLE_TYPE.PARAGRAPH)
        incident_title_paragraph_format = incident_title_style.paragraph_format
        incident_title_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        incident_title_font = incident_title_style.font
        incident_title_font.name = 'Times New Roman'
        incident_title_font.size = Pt(14)
        incident_title_font.bold = True

        title = self.document.add_heading('Отчет по сетевым инцидентам', level=1)
        title.style = title_style
    
        body_style = self.document.styles.add_style('Текст отчёта', WD_STYLE_TYPE.PARAGRAPH)
        body_paragraph_format = body_style.paragraph_format
        body_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        body_font = body_style.font
        body_font.name = 'Times New Roman'
        body_font.size = Pt(14)
        body_style.base_style = self.document.styles['Normal']

        for incident in self.incidents: # ИСПРАВИТЬ
            incident_title = self.document.add_heading(f'Инцидент №{incident.id}', level=2)
            incident_title.style = incident_title_style
            self.document.add_paragraph(f'Дата и время: {incident.timestamp}', style=body_style)
            self.document.add_paragraph(f'Тип атаки: {incident.attack_type}', style=body_style)
            #self.document.add_paragraph(f'Уровень угрозы: {incident.threat_level}',style=body_style)
            self.document.add_paragraph(f'Описание: {self.description[incident.attack_type]}',style=body_style)
            self.document.add_paragraph('')  # Добавляем пустой абзац между инцидентами
            self.add_packet_table(incident.packets)
            #self.add_packet_info_table(incident.packets)
            self.add_recommendations(incident.attack_type)

    def add_packet_table(self, packet_data):
        # Создаем стиль для заголовка таблицы
        table_heading_style = self.document.styles['Heading 1']
        table_heading_font = table_heading_style.font
        table_heading_font.name = 'Times New Roman'
        table_heading_font.size = Pt(12)
        table_heading_font.color.rgb = RGBColor(0,0,0)
        table_heading_font.bold = True

        # Создаем стиль для содержимого таблицы
        table_content_style = self.document.styles['Normal']
        table_content_font = table_content_style.font
        table_content_font.name = 'Times New Roman'
        table_content_font.size = Pt(8)

        # Добавляем заголовок таблицы
        table_heading = self.document.add_heading('Таблица захваченных IP пакетов', level=1)
        table_heading.style = table_heading_style

        # Создаем таблицу
        table = self.document.add_table(rows=1, cols=7)
        table.autofit = False  # Отключаем автонастройку размеров ячеек

        # Задаем ширину колонок
        column_widths = [0.5, 3, 3.5, 3.5, 3, 1.5, 4]
        for i, width in enumerate(column_widths):
            table.columns[i].width = Pt(width)

        # Добавляем заголовки столбцов
        column_headers = ['ID', 'Время', 'IP-источника', 'IP-назначения', 'Протокол', 'Длина', 'Информация']
        for i, header in enumerate(column_headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].style = table_heading_style

        # Добавляем данные в таблицу
        for packet in packet_data:
            row = table.add_row()
            print(f"{packet.id} {packet.time} {packet.source} {packet.destination} {packet.proto} {packet.length} {packet.raw_info}")
            print(len(row.cells))
            row.cells[0].text = str(packet.id)
            row.cells[1].text = packet.time
            row.cells[2].text = packet.source
            row.cells[3].text = packet.destination
            row.cells[4].text = packet.proto
            row.cells[5].text = str(packet.length)
            row.cells[6].text = packet.raw_info
            
        # Добавляем информацию о пакете в ячейку с использованием стилей
        #    info_cell = row.cells[7]
        #    info_cell.text = ''  # Очищаем ячейку перед заполнением
        #    for info_item in packet.info:
        #        info_cell.add_paragraph(info_item["header"], style=table_content_style)
        #        info_cell.add_paragraph(info_item["info"], style=table_content_style)
        table.style = 'Table Grid'
    
    def add_packet_info_table(self, packet_data):
        # Создаем стиль для заголовка таблицы
        table_heading_style = self.document.styles['Heading 1']
        table_heading_font = table_heading_style.font
        table_heading_font.name = 'Times New Roman'
        table_heading_font.size = Pt(12)
        table_heading_font.bold = True

        # Создаем стиль для содержимого таблицы
        table_content_style = self.document.styles['Normal']
        table_content_font = table_content_style.font
        table_content_font.name = 'Times New Roman'
        table_content_font.size = Pt(10)

        # Добавляем заголовок таблицы
        table_heading = self.document.add_heading('Информация захваченных IP пакетов', level=1)
        table_heading.style = table_heading_style

        # Создаем таблицу
        table = self.document.add_table(rows=1, cols=6)
        table.autofit = False  # Отключаем автонастройку размеров ячеек

        # Задаем ширину колонок
        column_widths = [1.5, 2, 2, 2, 2, 2]
        for i, width in enumerate(column_widths):
            table.columns[i].width = Pt(width)

        # Добавляем заголовки столбцов
        column_headers = ['ID', '1-й слой', '2-й слой', '3-й слой', '4-й слой', '5-й слой']
        for i, header in enumerate(column_headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].style = table_heading_style

        # Добавляем данные в таблицу
        for packet in packet_data:
            row = table.add_row()
            #print(f"{packet.id} {packet.time} {packet.source} {packet.destination} {packet.proto} {packet.length} {packet.raw_info}")
            #print(len(row.cells))
            row.cells[0].text = str(packet.id)
            row.cells[1].text = packet.info[0]["info"]
            row.cells[2].text = packet.info[1]["info"]
            row.cells[3].text = ""
            row.cells[4].text = ""
            row.cells[5].text = ""

            
        # Добавляем информацию о пакете в ячейку с использованием стилей
        #    info_cell = row.cells[7]
        #    info_cell.text = ''  # Очищаем ячейку перед заполнением
        #    for info_item in packet.info:
        #        info_cell.add_paragraph(info_item["header"], style=table_content_style)
        #        info_cell.add_paragraph(info_item["info"], style=table_content_style)
        table.style = 'Table Grid'

    def add_recommendations(self, attack_type):
        # Создаем стиль для заголовка рекомендаций
        recommendations_title_style = self.document.styles.add_style('Заголовок рекоммендаций', WD_STYLE_TYPE.PARAGRAPH)
        recommendations_title_paragraph_format = recommendations_title_style.paragraph_format
        recommendations_title_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        recommendations_title_font = recommendations_title_style.font
        recommendations_title_font.name = 'Times New Roman'
        recommendations_title_font.size = Pt(14)
        recommendations_title_font.bold = True

        # Добавляем заголовок рекомендаций
        recommendations_heading = self.document.add_heading('Рекомендации по устранению', level=1)
        recommendations_heading.style = recommendations_title_style
        
        

        # Добавляем рекомендации
        for recommendation in self.recommendations[attack_type]:
            body_style = self.document.styles.add_style('Текст рекоммендаций', WD_STYLE_TYPE.PARAGRAPH)
            body_paragraph_format = body_style.paragraph_format
            body_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            body_font = body_style.font
            body_font.name = 'Times New Roman'
            body_font.size = Pt(14)
            body_style.base_style = self.document.styles['Normal']
            paragraph = self.document.add_paragraph(recommendation, style=body_style)


    def generate_report(self, doc_path):
        # Добавляем информацию об атаке
        self.add_attack_info()
        # Сохраняем документ
        self.document.save(doc_path)


class DatePickWidget(ft.UserControl):
    def __init__(self, master):
        super().__init__()
        self.master = master
        self.datepicker = None
        self.holidays = [datetime(2023, 4, 25), datetime(2023, 5, 1), datetime(2023, 6, 2)]
        self.selected_locale = "ru_RU"
        

        self.dlg_modal = ft.AlertDialog(
            modal=True,
            title=ft.Text("Выбрать дату"),
            actions=[
                ft.Row([
                    ft.Row([
                        # ft.IconButton(icon=ft.icons.SETTINGS_OUTLINED),
                        ]),
                    ft.Row([ft.TextButton("Отмена", on_click=self.cancel_dlg),ft.TextButton("Подтвердить", on_click=self.confirm_dlg),]),
                ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN)
            ],
            actions_alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
            actions_padding=5,
            content_padding=0
        )

        self.tf = ft.TextField(label="Выбрать дату", dense=True, hint_text="дд.мм.гггг", width=260, height=40, on_change=self.master.singledatesearch)
        self.cal_ico = ft.TextButton(
            icon=ft.icons.CALENDAR_MONTH, 
            on_click=self.open_dlg_modal, 
            height=40,
            width=40,
            right=0,
            style=ft.ButtonStyle(
                padding=ft.Padding(4,0,0,0),
                shape={
                        ft.MaterialState.DEFAULT: ft.RoundedRectangleBorder(radius=1),
                    },
            ))

        self.st = ft.Stack(
            [
                self.tf,
                self.cal_ico,
            ]
        )
        self.from_to_text = ft.Text(visible=False)

    def build(self):
        return ft.Column(
            [
            self.st,
            self.from_to_text
            ]
        )
    def confirm_dlg(self, e):
        self.master.rangedatesearch()
        self.dlg_modal.open = False
        self.update()
        self.page.update()
    
    def update_result(self, selected_data):
        if len(selected_data) > 1:
            self.from_to_text.value = f"С: {selected_data[0].strftime('%d.%m.%Y')} по: {selected_data[1].strftime('%d.%m.%Y')}"
            self.from_to_text.visible = True
            self.master.start_date = selected_data[0].strftime('%d.%m.%Y')
            self.master.end_date = selected_data[1].strftime('%d.%m.%Y')    
    
    def cancel_dlg(self, e):
        self.dlg_modal.open = False
        self.page.update()

    def open_dlg_modal(self, e):
        self.tf.value = None
        self.datepicker = DatePicker(
            selection_type = 2,
            selected_date=[self.tf.value] if self.tf.value else None,
            holidays=self.holidays,
            locale=self.selected_locale,
            on_change=self.update_result
            )
        self.page.dialog = self.dlg_modal
        self.dlg_modal.content = self.datepicker
        self.dlg_modal.open = True
        self.page.update()

    def _to_datetime(self, date_str=None):
        if date_str:
            return datetime.strptime(date_str, "%d-%m-%Y %H:%M:%S")
        else:
            return None

class MyCheckBox(ft.UserControl):
    def __init__(self, master, id:int):
        self.master = master
        self.id = id        
        self.cb = ft.Checkbox(value=True, on_change=self.on_change)
        super().__init__()

    def on_change(self, e):
        if self.cb.value:            
            self.master.enabled_log_list.append(self.id)
            if len(self.master.enabled_log_list) == (len(self.master.attack_logs_data)):
                self.master.data_header.controls[0].content.value = True
                self.master.data_header.controls[0].content.update()
        else:
            self.master.data_header.controls[0].content.value = False
            self.master.data_header.controls[0].content.update()
            self.master.enabled_log_list.remove(self.id)
        print(self.master.enabled_log_list)

    def update_view(self):
        self.cb.update()

    def build(self):
        return self.cb

class LogPageClass(ft.UserControl):
    def __init__(self, master, attack_logs:list):
        self.master = master
        self.attack_logs_data = attack_logs
        #self.attack_logs_data = self.get_logs_data()
        self.enabled_log_list = []
        self.datepick_widget = DatePickWidget(self)
        self.start_date = None
        self.end_date = None
        self.cur_date = None
        self.cur_filter = None
        self.data_header = ft.Row(
            controls=[
                ft.Container(
                    expand=36,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),                    
                    content=ft.Checkbox(label="Выбрать все", value=True, on_change=self.choose_all)
                ),
                ft.Container(
                    expand=28,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),                    
                    content=ft.Text("ID атаки", text_align="center")
                ),
             ft.Container(
                    expand=58,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),
                    content=ft.Text("Класс атаки", text_align="center"),
                ),
             ft.Container(
                    expand=51,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),
                    content=ft.Text("Время атаки", text_align="center"),
                ),
            ft.Container(
                    expand=51,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),
                    content=ft.Text("Источник атаки", text_align="center"),
                ),             
            ],
            alignment=ft.alignment.center,
            spacing=0            
            #expand=True
        )
        self.logs_table = ft.DataTable(
        columns=[
            ft.DataColumn(ft.Text("")),
            ft.DataColumn(ft.Text("")),
            ft.DataColumn(ft.Text("")),
            ft.DataColumn(ft.Text("")),
            ft.DataColumn(ft.Text("")),
        ],
        rows=[],
        expand=True,
        vertical_lines=ft.border.BorderSide(1, "red"),
        heading_row_height=0,
        )
        for i in range(len(self.attack_logs_data)):
            self.logs_table.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(MyCheckBox(self, self.attack_logs_data[i].id)),
                        ft.DataCell(ft.Text(self.attack_logs_data[i].id)),
                        ft.DataCell(ft.Text(self.attack_logs_data[i].attack_type)),
                        ft.DataCell(ft.Text(self.attack_logs_data[i].timestamp)),
                        ft.DataCell(ft.Text(self.attack_logs_data[i].source))
                    ]
                )                
            )
            self.enabled_log_list.append(i)

        self.scroll_logs_table = ft.ListView(expand=1, spacing=10, padding=20)
        self.scroll_logs_table.controls.append(self.logs_table)
        self.nameinput = ft.TextField(label="Найти запись",
            on_change=self.inputsearch, height=40, border_color=ft.colors.RED_600
        )
        self.datanotfound = ft.Row([ft.Text("Запись не найдена",
                weight="bold",
                size=20,
                text_align="center",
        )], alignment=ft.MainAxisAlignment.SPACE_AROUND)
        self.datanotfound.visible = False
        super().__init__()

    def get_logs_data(self):
        try:
        # Чтение конфигурационных данных из файла .env
            host = config('DATABASE_HOST')
            port = config('DATABASE_PORT')
            database = config('DATABASE_NAME')
            user = config('DATABASE_USER')
            password = config('DATABASE_PASSWORD')
            table_name = config('DATABASE_TABLE')

            connection_string = f"postgresql+psycopg2://{user}:{password}@{host}:{port}/{database}"
            engine = create_engine(connection_string)

            query = f"SELECT * FROM {table_name};"
            df = pd.read_sql_query(query, engine)
            engine.dispose()
            return df
        except Exception as e:
            print(f"Ошибка при подключении к базе данных: {e}")
            return None

    def rangefilter(self, date):
        timestamp_str = date.timestamp
        input_datetime = datetime.strptime(timestamp_str, "%d.%m.%Y %H:%M:%S")
        start_datetime = datetime.strptime(self.start_date, "%d.%m.%Y")
        end_datetime = datetime.strptime(self.end_date, "%d.%m.%Y")

        if start_datetime <= input_datetime <= end_datetime:
            return True
        else:
            return False

    def singlefilter(self, date):
        timestamp_str = date.timestamp
        input_datetime = str(datetime.strptime(timestamp_str, "%d.%m.%Y %H:%M:%S")).split(' ')[0]
        if input_datetime == str(self.cur_date).split(' ')[0]:
            return True
        else:
            return False

    def rangedatesearch(self):
        if self.cur_filter == None:
            myfilter = list(filter(self.rangefilter, self.attack_logs_data))
            self.cur_filter = myfilter
        else:
            myfilter = list(filter(self.rangefilter, self.cur_filter))        
        self.logs_table.rows = []

        if len(myfilter) > 0:
            self.datanotfound.visible = False
            for i in range(len(myfilter)):
                self.logs_table.rows.append(
                    ft.DataRow(
                        cells=[
                        ft.DataCell(MyCheckBox(self, myfilter[i].id)),
                        ft.DataCell(ft.Text(myfilter[i].id)),
                        ft.DataCell(ft.Text(myfilter[i].attack_type)),
                        ft.DataCell(ft.Text(myfilter[i].timestamp)),
                        ft.DataCell(ft.Text(myfilter[i].source))
                        ]
                    ) 
                )
                self.enabled_log_list.clear()
                self.enabled_log_list.append(i)
                self.update()
        else:              
            self.datanotfound.visible = True
            self.update()
        
    def singledatesearch(self, e):
        try:
            self.cur_date = datetime.strptime(self.datepick_widget.tf.value, "%d.%m.%Y")
            myfilter = list(filter(self.singlefilter, self.attack_logs_data))        
            self.logs_table.rows = []
            if not self.datepick_widget.tf.value == "":
                self.datepick_widget.from_to_text.value = self.datepick_widget.tf.value
                self.datepick_widget.from_to_text.update()
                if len(myfilter) > 0:
                    self.datanotfound.visible = False
                    for i in range(len(myfilter)):
                        self.logs_table.rows.append(
                            ft.DataRow(
                                cells=[
                                ft.DataCell(MyCheckBox(self, myfilter[i].id)),
                                ft.DataCell(ft.Text(myfilter[i].id)),
                                ft.DataCell(ft.Text(myfilter[i].attack_type)),
                                ft.DataCell(ft.Text(myfilter[i].timestamp)),
                                ft.DataCell(ft.Text(myfilter[i].source))
                                ]
                            ) 
                        )
                        self.enabled_log_list.clear()
                        self.enabled_log_list.append(i)
                        self.update()
                else:                
                    self.datanotfound.visible = True
                    self.update()
            else:
                self.datanotfound.visible = False
                self.update()
                for i in range(len(self.attack_logs_data)):
                    self.logs_table.rows.append(
                        ft.DataRow(
                            cells=[
                                ft.DataCell(MyCheckBox(self, self.attack_logs_data[i].id)),
                                ft.DataCell(ft.Text(self.attack_logs_data[i].id)),
                                ft.DataCell(ft.Text(self.attack_logs_data[i].attack_type)),
                                ft.DataCell(ft.Text(self.attack_logs_data[i].timestamp)),
                                ft.DataCell(ft.Text(self.attack_logs_data[i].source))
                            ]
                        )
                    )
                    self.enabled_log_list.clear()
                    self.enabled_log_list.append(i)
                self.update()
        except ValueError:
            pass   

    def inputsearch(self, e):
        search_name = self.nameinput.value.lower()
        if self.cur_filter == None:
            myfilter = list(filter(lambda x:search_name in x.attack_type.lower(), self.attack_logs_data))
            self.cur_filter = myfilter
        else:
            myfilter = list(filter(lambda x:search_name in x.attack_type.lower(), self.cur_filter))
        self.logs_table.rows = []
        if not self.nameinput.value == "":
            self.update()
            if len(myfilter) > 0:
                self.datanotfound.visible = False
                for i in range(len(myfilter)):
                    self.logs_table.rows.append(
                       ft.DataRow(
                            cells=[
                            ft.DataCell(MyCheckBox(self, myfilter[i].id)),
                                ft.DataCell(ft.Text(myfilter[i].id)),
                                ft.DataCell(ft.Text(myfilter[i].attack_type)),
                                ft.DataCell(ft.Text(myfilter[i].timestamp)),
                                ft.DataCell(ft.Text(myfilter[i].source))
                            ]
                        ) 
                    )
                    self.enabled_log_list.clear()
                    self.enabled_log_list.append(i)
                    self.update()
            else:                
                self.datanotfound.visible = True
                self.update()
        else:
            self.cur_filter = None
            self.datanotfound.visible = False
            self.update()
            for i in range(len(self.attack_logs_data)):
                self.logs_table.rows.append(
                    ft.DataRow(
                        cells=[
                            ft.DataCell(MyCheckBox(self, self.attack_logs_data[i].id)),
                            ft.DataCell(ft.Text(self.attack_logs_data[i].id)),
                            ft.DataCell(ft.Text(self.attack_logs_data[i].attack_type)),
                            ft.DataCell(ft.Text(self.attack_logs_data[i].timestamp)),
                            ft.DataCell(ft.Text(self.attack_logs_data[i].source))
                        ]
                    )
                )
                self.enabled_log_list.clear()
                self.enabled_log_list.append(i)
            self.update()    
    
    def choose_all(self, e):
        self.enabled_log_list.clear()
        if self.data_header.controls[0].content.value:            
            for i in range(len(self.logs_table.rows)):
                self.logs_table.rows[i].cells[0].content.cb.value = True
                self.logs_table.rows[i].cells[0].content.update_view()
                self.enabled_log_list.append(self.logs_table.rows[i].cells[0].content.id)                
        else:
            for i in range(len(self.logs_table.rows)):
                self.logs_table.rows[i].cells[0].content.cb.value = False
                self.logs_table.rows[i].cells[0].content.update_view()
    
    def save_report(self, path):
        report = AttackReport(self.attack_logs_data, self.enabled_log_list)
        report.generate_report(path)

    def mysavefile(self, e:ft.FilePickerResultEvent):
        save_location = e.path
        try:
          if save_location:
            self.save_report(save_location+".docx")
            self.master.open_dlg_modal()
        except Exception as e:
          print("Ошибка", e)

    def build(self):
        return ft.Column(
            controls=[
                ft.Row(
                    controls=[
                        ft.Container(
                    height=50,
                    padding=5,                    
                    alignment=ft.alignment.center,
                    margin=15,
                    content=self.nameinput
                    ),
                    self.datepick_widget,
                    ft.OutlinedButton("Сформировать отчёт", on_click=lambda _: self.master.saveme.save_file(dialog_title="Сохранить как",
                    file_type=ft.FilePickerFileType.CUSTOM,
                    allowed_extensions = ["docx"]))
                    ]
                ),                
                
                ft.Container(
                    expand=True,
                    border_radius=10,
                    margin=15,
                    padding=0,
                    alignment=ft.alignment.center,                   
                    #bgcolor=ft.colors.GREY,
                    border=ft.border.all(1, ft.colors.RED_600),
                    content=ft.Column(
                        controls=[
                            ft.Container(
                                expand=1,
                                padding=0,
                                alignment=ft.alignment.center,
                                border=ft.border.all(1, ft.colors.RED_600),
                                content=self.data_header                                
                            ),                            
                            self.datanotfound,                            
                            ft.Container(
                                expand=10,
                                padding=0,
                                alignment=ft.alignment.center,
                                border=ft.border.all(1, ft.colors.RED_600),
                                content=self.scroll_logs_table
                            ),                           
                            
                        ]
                    )
                )
            ]
        )

def BuildLogPage(master, animation_style, attack_logs):          
    return ft.Container(
        alignment=ft.alignment.center,
        offset=ft.transform.Offset(0,0),
        animate_offset=animation_style,
        bgcolor="#1f2128",
        #bgcolor="white",
        content=LogPageClass(master, attack_logs)
        )
