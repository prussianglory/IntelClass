import flet as ft
from datepicker.datepicker import DatePicker
from datepicker.selection_type import SelectionType
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL

class AdminReportBuilder:
    def __init__(self, filepicker, admin_logs:list, logs_id_list:list):
        self.filepicker = filepicker
        self.admin_logs = admin_logs
        self.logs_id_list = logs_id_list
        self.description = {
            "Worms":"",
            "DoS":"",
            "Exploits":"",
            "Analysis":"",
            "Backdoor":"",
            "Shellcode":""
        }
        self.recommendations = {
            "Worms":"",
            "DoS":"",
            "Exploits":"",
            "Analysis":"",
            "Backdoor":"",
            "Shellcode":""
        }
    
    def create_packet_table(doc_path, packet_data):
        document = Document()

        table_heading_style = document.styles['Heading 1']
        table_heading_font = table_heading_style.font
        table_heading_font.name = 'Times New Roman'
        table_heading_font.size = Pt(12)
        table_heading_font.bold = True

        # Создаем стиль для содержимого таблицы
        table_content_style = document.styles['Normal']
        table_content_font = table_content_style.font
        table_content_font.name = 'Arial'
        table_content_font.size = Pt(10)

        # Добавляем заголовок таблицы
        table_heading = document.add_heading('Таблица захваченных IP пакетов', level=1)
        table_heading.style = table_heading_style

        # Создаем таблицу
        table = document.add_table(rows=1, cols=4)
        table.autofit = False  # Отключаем автонастройку размеров ячеек

        # Задаем ширину колонок
        column_widths = [1.5, 2, 2, 2]
        for i, width in enumerate(column_widths):
            table.columns[i].width = Pt(width)

        # Добавляем заголовки столбцов
        column_headers = ['ID', 'Класс события', 'Время', 'Успешность попытки']
        for i, header in enumerate(column_headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].style = table_heading_style

        # Добавляем данные в таблицу
        for packet in packet_data:
            row = table.add_row()
            for i, header in enumerate(column_headers):
                cell = row.cells[i]
                cell.text = str(packet.get(header, ''))  # Получаем значение из словаря
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].style = table_content_style
        table.style = 'Table Grid'

    def create_report(self, doc_path):
        document = Document()

        title_style = document.styles.add_style('Заголовок отчёта', WD_STYLE_TYPE.PARAGRAPH)
        title_paragraph_format = title_style.paragraph_format
        title_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(16)
        title_font.bold = True

        incident_title_style = document.styles.add_style('Заголовок инцидента', WD_STYLE_TYPE.PARAGRAPH)
        incident_title_paragraph_format = incident_title_style.paragraph_format
        incident_title_paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        incident_title_font = incident_title_style.font
        incident_title_font.name = 'Times New Roman'
        incident_title_font.size = Pt(14)
        incident_title_font.bold = True

        title = document.add_heading('Отчет по сетевым инцидентам', level=1)
        title.style = title_style
    
        body_style = document.styles.add_style('Текст отчёта', WD_STYLE_TYPE.PARAGRAPH)
        body_paragraph_format = body_style.paragraph_format
        body_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        body_font = body_style.font
        body_font.name = 'Times New Roman'
        body_font.size = Pt(14)
        body_style.base_style = document.styles['Normal']

        incidents = [
        {
            'number': 1,
            'type': 'SQL инъекция',
            'date_time': '2023-12-15, 14:30',
            'threat_level': 'Высокий',
            'description': 'Неопознанный злоумышленник осуществил SQL-инъекцию...'
        },
        {
            'number': 2,
            'type': 'DDoS-атака',
            'date_time': '2023-12-17, 09:45',
            'threat_level': 'Средний',
            'description': 'Серверы были перегружены огромным объемом фальшивых запросов...'
        },
        {
            'number': 3,
            'type': 'Фишинговая атака',
            'date_time': '2023-12-18, 11:15',
            'threat_level': 'Низкий',
            'description': 'Пользователи получили электронные письма, представляющиеся как официальные сообщения...'
        },
        ]

        for incident in incidents:
            incident_title = document.add_heading(f'Инцидент №{incident["number"]}: {incident["type"]}', level=2)
            incident_title.style = incident_title_style
            document.add_paragraph(f'Дата и время: {incident["date_time"]}', style=body_style)
            document.add_paragraph(f'Тип атаки: {incident["type"]}', style=body_style)
            document.add_paragraph(f'Уровень угрозы: {incident["threat_level"]}',style=body_style)
            document.add_paragraph(f'Описание: {incident["description"]}',style=body_style)
            document.add_paragraph('')  # Добавляем пустой абзац между инцидентами

        # Сохраняем документ
        document.save(doc_path)


class DatePickWidget(ft.UserControl):
    def __init__(self, master):
        super().__init__()
        self.master = master
        self.datepicker = None
        self.holidays = [datetime(2023, 4, 25), datetime(2023, 5, 1), datetime(2023, 6, 2)]
        self.selected_locale = "ru_RU"
        

        self.dlg_modal = ft.AlertDialog(
            modal=True,
            title=ft.Text("Выбрать дату"),
            actions=[
                ft.Row([
                    ft.Row([
                        # ft.IconButton(icon=ft.icons.SETTINGS_OUTLINED),
                        ]),
                    ft.Row([ft.TextButton("Отмена", on_click=self.cancel_dlg),ft.TextButton("Подтвердить", on_click=self.confirm_dlg),]),
                ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN)
            ],
            actions_alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
            actions_padding=5,
            content_padding=0
        )

        self.tf = ft.TextField(label="Выбрать дату", dense=True, hint_text="дд.мм.гггг", width=260, height=40, on_change=self.master.singledatesearch)
        self.cal_ico = ft.TextButton(
            icon=ft.icons.CALENDAR_MONTH, 
            on_click=self.open_dlg_modal, 
            height=40,
            width=40,
            right=0,
            style=ft.ButtonStyle(
                padding=ft.Padding(4,0,0,0),
                shape={
                        ft.MaterialState.DEFAULT: ft.RoundedRectangleBorder(radius=1),
                    },
            ))

        self.st = ft.Stack(
            [
                self.tf,
                self.cal_ico,
            ]
        )
        self.from_to_text = ft.Text(visible=False)

    def build(self):
        return ft.Column(
            [
            self.st,
            self.from_to_text
            ]
        )
    def confirm_dlg(self, e):
        self.master.rangedatesearch()
        self.dlg_modal.open = False
        self.update()
        self.page.update()
    
    def update_result(self, selected_data):
        if len(selected_data) > 1:
            self.from_to_text.value = f"С: {selected_data[0].strftime('%d.%m.%Y')} по: {selected_data[1].strftime('%d.%m.%Y')}"
            self.from_to_text.visible = True
            self.master.start_date = selected_data[0].strftime('%d.%m.%Y')
            self.master.end_date = selected_data[1].strftime('%d.%m.%Y')    
    
    def cancel_dlg(self, e):
        self.dlg_modal.open = False
        self.page.update()

    def open_dlg_modal(self, e):
        self.tf.value = None
        self.datepicker = DatePicker(
            selection_type = 2,
            selected_date=[self.tf.value] if self.tf.value else None,
            holidays=self.holidays,
            locale=self.selected_locale,
            on_change=self.update_result
            )
        self.page.dialog = self.dlg_modal
        self.dlg_modal.content = self.datepicker
        self.dlg_modal.open = True
        self.page.update()

    def _to_datetime(self, date_str=None):
        if date_str:
            return datetime.strptime(date_str, "%d-%m-%Y %H:%M:%S")
        else:
            return None

class MyCheckBox(ft.UserControl):
    def __init__(self, master, id:int):
        self.master = master
        self.id = id        
        self.cb = ft.Checkbox(value=True, on_change=self.on_change)
        super().__init__()

    def on_change(self, e):
        if self.cb.value:            
            self.master.enabled_log_list.append(self.id)
            if len(self.master.enabled_log_list) == (len(self.master.admin_logs_data)):
                self.master.data_header.controls[0].content.value = True
                self.master.data_header.controls[0].content.update()
        else:
            self.master.data_header.controls[0].content.value = False
            self.master.data_header.controls[0].content.update()
            self.master.enabled_log_list.remove(self.id)
        print(self.master.enabled_log_list)

    def update_view(self):
        self.cb.update()

    def build(self):
        return self.cb

class AdminLogPageClass(ft.UserControl):
    def __init__(self, master, admin_logs:list):
        self.master = master
        self.admin_logs_data = admin_logs
        self.enabled_log_list = []
        self.datepick_widget = DatePickWidget(self)
        self.start_date = None
        self.end_date = None
        self.cur_date = None
        self.data_header = ft.Row(
            controls=[
                ft.Container(
                    expand=32,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),                    
                    content=ft.Checkbox(label="Выбрать все", value=True, on_change=self.choose_all)
                ),
                ft.Container(
                    expand=24,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),                    
                    content=ft.Text("ID события", text_align="center")
                ),
             ft.Container(
                    expand=80,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),
                    content=ft.Text("Класс события", text_align="center"),
                ),
             ft.Container(
                    expand=44,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),
                    content=ft.Text("Время события", text_align="center"),
                ),
                ft.Container(
                    expand=41,                  
                    alignment=ft.alignment.center,
                    border=ft.border.all(1, ft.colors.RED_600),
                    content=ft.Text("Успешность попытки", text_align="center"),
                ),            
            ],
            alignment=ft.alignment.center,
            spacing=0            
            #expand=True
        )
        self.logs_table = ft.DataTable(
        columns=[
            ft.DataColumn(ft.Text("")),
            ft.DataColumn(ft.Text("")),
            ft.DataColumn(ft.Text("")),
            ft.DataColumn(ft.Text("")),
            ft.DataColumn(ft.Text("")),
        ],
        rows=[],
        expand=True,
        vertical_lines=ft.border.BorderSide(1, "red"),
        heading_row_height=0,
        )
        for i in range(len(self.admin_logs_data)):
            self.logs_table.rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(MyCheckBox(self, i)),
                        ft.DataCell(ft.Text(self.admin_logs_data[i].id)),
                        ft.DataCell(ft.Text(self.admin_logs_data[i].log_class)),
                        ft.DataCell(ft.Text(self.admin_logs_data[i].ts)),
                        ft.DataCell(ft.Text(self.admin_logs_data[i].success))
                    ]
                )                
            )
            self.enabled_log_list.append(i)

        self.scroll_logs_table = ft.ListView(expand=1, spacing=10, padding=20)
        self.scroll_logs_table.controls.append(self.logs_table)
        self.nameinput = ft.TextField(label="Найти запись",
            on_change=self.inputsearch, height=40, border_color=ft.colors.RED_600
        )
        self.datanotfound = ft.Row([ft.Text("Запись не найдена",
                weight="bold",
                size=20,
                text_align="center",
        )], alignment=ft.MainAxisAlignment.SPACE_AROUND)
        self.datanotfound.visible = False
        super().__init__()

    def rangefilter(self, date):
        timestamp_str = date.timestamp
        input_datetime = datetime.strptime(timestamp_str, "%d.%m.%Y %H:%M:%S")
        start_datetime = datetime.strptime(self.start_date, "%d.%m.%Y")
        end_datetime = datetime.strptime(self.end_date, "%d.%m.%Y")

        if start_datetime <= input_datetime <= end_datetime:
            return True
        else:
            return False

    def singlefilter(self, date):
        timestamp_str = date.timestamp
        input_datetime = str(datetime.strptime(timestamp_str, "%d.%m.%Y %H:%M:%S")).split(' ')[0]
        if input_datetime == str(self.cur_date).split(' ')[0]:
            return True
        else:
            return False


    def rangedatesearch(self):
        myfilter = list(filter(self.rangefilter, self.admin_logs_data))        
        self.logs_table.rows = []
        if len(myfilter) > 0:
            self.datanotfound.visible = False
            for i in range(len(myfilter)):
                self.logs_table.rows.append(
                    ft.DataRow(
                        cells=[
                        ft.DataCell(MyCheckBox(self, i)),
                        ft.DataCell(ft.Text(myfilter[i].id)),
                        ft.DataCell(ft.Text(myfilter[i].log_class)),
                        ft.DataCell(ft.Text(myfilter[i].ts)),
                        ft.DataCell(ft.Text(myfilter[i].success))
                        ]
                    ) 
                )
                self.enabled_log_list.clear()
                self.enabled_log_list.append(i)
                self.update()
        else:                
            self.datanotfound.visible = True
            self.update()
        
    def singledatesearch(self, e):
        try:
            self.cur_date = datetime.strptime(self.datepick_widget.tf.value, "%d.%m.%Y")
            myfilter = list(filter(self.singlefilter, self.admin_logs_data))        
            self.logs_table.rows = []
            if not self.datepick_widget.tf.value == "":
                if len(myfilter) > 0:
                    self.datanotfound.visible = False
                    for i in range(len(myfilter)):
                        self.logs_table.rows.append(
                            ft.DataRow(
                                cells=[
                                ft.DataCell(MyCheckBox(self, i)),
                                ft.DataCell(ft.Text(myfilter[i].id)),
                                ft.DataCell(ft.Text(myfilter[i].log_class)),
                                ft.DataCell(ft.Text(myfilter[i].ts)),
                                ft.DataCell(ft.Text(myfilter[i].success))
                                ]
                            ) 
                        )
                        self.enabled_log_list.clear()
                        self.enabled_log_list.append(i)
                        self.update()
                else:                
                    self.datanotfound.visible = True
                    self.update()
            else:
                self.datanotfound.visible = False
                self.update()
                for i in range(len(self.admin_logs_data)):
                    self.logs_table.rows.append(
                        ft.DataRow(
                            cells=[
                                ft.DataCell(MyCheckBox(self, i)),
                                ft.DataCell(ft.Text(self.admin_logs_data[i].id)),
                                ft.DataCell(ft.Text(self.admin_logs_data[i].log_class)),
                                ft.DataCell(ft.Text(self.admin_logs_data[i].ts)),
                                ft.DataCell(ft.Text(self.admin_logs_data[i].success))
                            ]
                        )
                    )
                    self.enabled_log_list.clear()
                    self.enabled_log_list.append(i)
                self.update()
        except ValueError:
            pass   

    def inputsearch(self, e):
        search_name = self.nameinput.value.lower()
        myfilter = list(filter(lambda x:search_name in x.log_class.lower(), self.admin_logs_data))        
        self.logs_table.rows = []
        if not self.nameinput.value == "":
            if len(myfilter) > 0:
                self.datanotfound.visible = False
                for i in range(len(myfilter)):
                    self.logs_table.rows.append(
                       ft.DataRow(
                            cells=[
                            ft.DataCell(MyCheckBox(self, i)),
                                ft.DataCell(ft.Text(myfilter[i].id)),
                                ft.DataCell(ft.Text(myfilter[i].log_class)),
                                ft.DataCell(ft.Text(myfilter[i].ts)),
                                ft.DataCell(ft.Text(myfilter[i].success))
                            ]
                        ) 
                    )
                    self.enabled_log_list.clear()
                    self.enabled_log_list.append(i)
                    self.update()
            else:                
                self.datanotfound.visible = True
                self.update()
        else:
            self.datanotfound.visible = False
            self.update()
            for i in range(len(self.admin_logs_data)):
                self.logs_table.rows.append(
                    ft.DataRow(
                        cells=[
                            ft.DataCell(MyCheckBox(self, i)),
                            ft.DataCell(ft.Text(self.admin_logs_data[i].id)),
                            ft.DataCell(ft.Text(self.admin_logs_data[i].log_class)),
                            ft.DataCell(ft.Text(self.admin_logs_data[i].ts)),
                            ft.DataCell(ft.Text(self.admin_logs_data[i].success)),
                        ]
                    )
                )
                self.enabled_log_list.clear()
                self.enabled_log_list.append(i)
            self.update()    
    
    def choose_all(self, e):
        self.enabled_log_list.clear()
        if self.data_header.controls[0].content.value:            
            for i in range(len(self.logs_table.rows)):
                self.logs_table.rows[i].cells[0].content.cb.value = True
                self.logs_table.rows[i].cells[0].content.update_view()
                self.enabled_log_list.append(self.logs_table.rows[i].cells[0].content.id)                
        else:
            for i in range(len(self.logs_table.rows)):
                self.logs_table.rows[i].cells[0].content.cb.value = False
                self.logs_table.rows[i].cells[0].content.update_view()

    def build(self):
        return ft.Column(
            controls=[
                ft.Row(
                    controls=[
                        ft.Container(
                    height=50,
                    padding=5,                    
                    alignment=ft.alignment.center,
                    margin=15,
                    content=self.nameinput
                    ),
                    self.datepick_widget,
                    ft.OutlinedButton("Сформировать отчёт")
                    ]
                ),                
                
                ft.Container(
                    expand=True,
                    border_radius=10,
                    margin=15,
                    padding=0,
                    alignment=ft.alignment.center,                   
                    #bgcolor=ft.colors.GREY,
                    border=ft.border.all(1, ft.colors.RED_600),
                    content=ft.Column(
                        controls=[
                            ft.Container(
                                expand=1,
                                padding=0,
                                alignment=ft.alignment.center,
                                border=ft.border.all(1, ft.colors.RED_600),
                                content=self.data_header                                
                            ),                            
                            self.datanotfound,                            
                            ft.Container(
                                expand=10,
                                padding=0,
                                alignment=ft.alignment.center,
                                border=ft.border.all(1, ft.colors.RED_600),
                                content=self.scroll_logs_table
                            ),                           
                            
                        ]
                    )
                )
            ]
        )

def BuildAdminLogPage(master, animation_style, admin_logs):          
    return ft.Container(
        alignment=ft.alignment.center,
        offset=ft.transform.Offset(0,0),
        animate_offset=animation_style,
        bgcolor="#1f2128",
        content=AdminLogPageClass(master, admin_logs)
        )
